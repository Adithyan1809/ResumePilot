import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    if level == 0:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_paragraph(doc, text):
    return doc.add_paragraph(text)

def create_doc():
    doc = Document()

    # Title
    add_heading(doc, 'ResumePilot: Comprehensive System Documentation', 0)

    doc.add_paragraph('ResumePilot has evolved from a basic resume builder into an Ultra-Premium AI Employability Operating System. This document outlines the entire journey, architectural decisions, database models, and the specialized AI pipeline that powers the platform.')

    # Phase 1
    add_heading(doc, '1. Project Evolution & Phased Implementation', 1)
    
    add_heading(doc, 'Phase 1: Foundation & Core Infrastructure', 2)
    p = doc.add_paragraph()
    p.add_run('Goal: ').bold = True
    p.add_run('Establish a robust, scalable full-stack environment.\n')
    p.add_run('Actions:').bold = True
    doc.add_paragraph('• Set up a decoupled architecture: Next.js (React) frontend and FastAPI (Python) backend.', style='List Bullet')
    doc.add_paragraph('• Configured PostgreSQL database with SQLAlchemy ORM and Alembic migrations.', style='List Bullet')
    doc.add_paragraph('• Implemented JWT-based secure user authentication (signup, login, session management).', style='List Bullet')
    doc.add_paragraph('• Established the foundational UI using Tailwind CSS.', style='List Bullet')

    # Phase 2
    add_heading(doc, 'Phase 2: Master Profile & Data Ingestion', 2)
    p = doc.add_paragraph()
    p.add_run('Goal: ').bold = True
    p.add_run('Allow users to establish a "Source of Truth" for their career history.\n')
    p.add_run('Actions:').bold = True
    doc.add_paragraph('• Designed the master ProfileMemory JSON schema to hold unstructured career data.', style='List Bullet')
    doc.add_paragraph('• Built the /api/resume/upload pipeline to parse PDF resumes using PyPDF2 and LLM extraction.', style='List Bullet')
    doc.add_paragraph('• Created the Dashboard UI to display the master profile status.', style='List Bullet')

    # Phase 3
    add_heading(doc, 'Phase 3: The Tailoring Engine (V1)', 2)
    p = doc.add_paragraph()
    p.add_run('Goal: ').bold = True
    p.add_run('Match candidate profiles to specific Job Descriptions (JDs).\n')
    p.add_run('Actions:').bold = True
    doc.add_paragraph("• Integrated xAI's Grok API for intelligent text generation.", style='List Bullet')
    doc.add_paragraph('• Built the JD extraction mechanism to identify required skills, tone, and key responsibilities.', style='List Bullet')
    doc.add_paragraph('• Implemented basic prompt engineering to rewrite resume bullets to align with the JD.', style='List Bullet')

    # Phase 4
    add_heading(doc, 'Phase 4: ATS Optimization & Document Generation', 2)
    p = doc.add_paragraph()
    p.add_run('Goal: ').bold = True
    p.add_run('Ensure resumes pass automated tracking systems and look professional.\n')
    p.add_run('Actions:').bold = True
    doc.add_paragraph('• Developed the ATS Scoring algorithm (TF-IDF keyword matching, format compliance).', style='List Bullet')
    doc.add_paragraph('• Implemented deterministic export pipelines (pdfkit for PDF, python-docx for Word).', style='List Bullet')

    # Phase 5
    add_heading(doc, 'Phase 5: Recruiter-Grade Intelligence & Cinematic Redesign (V4)', 2)
    p = doc.add_paragraph()
    p.add_run('Goal: ').bold = True
    p.add_run('Eliminate "AI-sounding" output and elevate the product to Apple/Stripe-level aesthetics.\n')
    p.add_run('Actions:').bold = True
    doc.add_paragraph('• Backend: Replaced monolithic prompt with a 9-stage Agentic Orchestration Pipeline.', style='List Bullet')
    doc.add_paragraph('• Frontend: Completely redesigned the UI with Framer Motion, Glassmorphism, Zinc palettes.', style='List Bullet')

    # Architecture
    add_heading(doc, '2. System Architecture', 1)
    doc.add_paragraph('ResumePilot uses a modern, decoupled architecture optimizing for both high-performance AI processing and buttery-smooth client rendering.')
    doc.add_paragraph('• Frontend: Client / Next.js (Cinematic UI, Tailor Studio, Auth Context)')
    doc.add_paragraph('• Backend: FastAPI Server (REST API, Workflow Orchestrator, ATS Scoring, Document Generation)')
    doc.add_paragraph('• Narrative Pipeline (The 9 Engines): Connects directly with xAI Grok API.')
    doc.add_paragraph('• Infrastructure: PostgreSQL DB, External LLM APIs.')

    # DB Schema
    add_heading(doc, '3. Database Schema', 1)
    doc.add_paragraph('The database is built on PostgreSQL using SQLAlchemy. It tracks users, their master profile memory, and the history of tailoring sessions.')
    doc.add_paragraph('• USERS: id, email, hashed_password, full_name, is_active, created_at')
    doc.add_paragraph('• MASTER_PROFILES: id, user_id (FK), parsed_data (Source of Truth), created_at, updated_at')
    doc.add_paragraph('• TAILORED_RESUMES: id, user_id (FK), job_title, company, job_description, final_resume_data, ats_scoring, template_style, created_at')

    # Engines
    add_heading(doc, '4. The 9-Stage Narrative Intelligence Pipeline', 1)
    doc.add_paragraph('To prevent the "AI-generated" feel, the tailoring process routes through 9 specialized Python engines in workflow_orchestrator.py:')
    doc.add_paragraph('1. Candidate Voice Engine: Preserves authentic communication style and maturity level.', style='List Number')
    doc.add_paragraph('2. Project Context Binding: Ensures technologies are only mentioned in context.', style='List Number')
    doc.add_paragraph('3. Engineering Identity Engine: Shapes the narrative to match the target role.', style='List Number')
    doc.add_paragraph('4. Resume Humanization: Strips out recognizable LLM jargon.', style='List Number')
    doc.add_paragraph('5. Experience Calibration: Validates claims match years of experience.', style='List Number')
    doc.add_paragraph('6. Project Ranker: Reorders bullets based on JD relevance.', style='List Number')
    doc.add_paragraph('7. Story Flow Engine: Optimizes vertical reading experience.', style='List Number')
    doc.add_paragraph('8. Metric Realism: Caps exaggerated metrics.', style='List Number')
    doc.add_paragraph('9. Recruiter Polish: Manages bullet density, breaks, and readability.', style='List Number')

    # Tech Stack
    add_heading(doc, '5. Technology Stack', 1)
    add_heading(doc, 'Frontend (Cinematic UI)', 2)
    doc.add_paragraph('Framework: Next.js 16 (App Router)\nStyling: Tailwind CSS v4\nAnimations: Framer Motion\nIcons: Lucide React\nState: React Context API')
    
    add_heading(doc, 'Backend (API & Orchestration)', 2)
    doc.add_paragraph('Framework: FastAPI (Python 3)\nDatabase: PostgreSQL\nORM: SQLAlchemy / Alembic\nAI: xAI Grok API\nDocument Generation: pdfkit, python-docx')

    add_heading(doc, '6. Premium UI/UX Design System Guidelines', 1)
    doc.add_paragraph('• Colors: Zinc/Graphite neutrals with Electric Indigo and Purple accents.', style='List Bullet')
    doc.add_paragraph('• Typography: Geist (Vercel\'s typeface) for a premium feel.', style='List Bullet')
    doc.add_paragraph('• Glassmorphism: Heavy use of backdrop-blur (16px - 32px), semi-transparent borders.', style='List Bullet')
    doc.add_paragraph('• Motion: Apple-style spring easing curves. Avoid bouncy or erratic animations.', style='List Bullet')

    doc.save('C:/Users/HP/Pictures/ResumeBuilder/ResumePilot_Documentation.docx')
    print("Docx generated successfully!")

if __name__ == '__main__':
    create_doc()
