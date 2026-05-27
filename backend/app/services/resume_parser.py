"""
Resume parsing service for PDF and DOCX documents.
"""

import logging
import re
from typing import Any, Dict, Tuple

from docx import Document
import pdfplumber
import logging

_LOGGER = logging.getLogger(__name__)

# Optional OCR imports (used only if available)
try:
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image
    _OCR_AVAILABLE = True
except Exception:
    _OCR_AVAILABLE = False

from app.utils.text import clean_text, extract_emails, extract_phones, extract_urls, heal_bullet_text

logger = logging.getLogger(__name__)

# Heuristic header mappings to classify section content
SECTION_PATTERNS = {
    "summary": r"\b(summary|profile|objective|professional summary|about me|career objective|executive summary)\b",
    "experience": r"\b(experience|work experience|employment|employment history|work history|professional experience|professional background|career history)\b",
    "education": r"\b(education|academic background|academic history|degrees|credentials|academic qualifications)\b",
    "skills": r"\b(skills|technical skills|core competencies|technologies|expertise|areas of expertise|professional skills|skills & expertise)\b",
    "certifications": r"\b(certifications|certification|licenses|courses|accreditations|professional certifications)\b",
    "projects": r"\b(projects|key projects|personal projects|selected projects|academic projects)\b",
    "languages": r"\b(languages|language|language proficiency)\b",
    "awards": r"\b(awards|honours|honors|achievements|distinctions|fellowships)\b",
    "publications": r"\b(publications|patents|research papers|articles)\b",
    "volunteer": r"\b(volunteer|volunteering|volunteer experience|community service|social work)\b",
}


async def parse_document(file_path: str, content_type: str) -> Tuple[str, Dict[str, Any]]:
    """Parse a resume file (PDF or DOCX), extracting raw text and segmenting it into sections.

    Args:
        file_path: Absolute path to the file on disk.
        content_type: MIME type of the file.

    Returns:
        A tuple of (raw_text, parsed_sections_dict).
    """
    logger.info(f"Parsing document: {file_path} ({content_type})")

    # 1. Extract text based on file format
    if "pdf" in content_type or file_path.lower().endswith(".pdf"):
        raw_text = _extract_pdf_text(file_path)
        # If pdfplumber returned no or very little text, try OCR fallback
        if not raw_text or len(raw_text.strip()) < 200:
            _LOGGER.info("PDF text extraction returned little text; attempting OCR fallback.")
            ocr_text = _extract_pdf_text_with_ocr(file_path)
            if ocr_text and len(ocr_text.strip()) > len(raw_text.strip()):
                raw_text = (raw_text or "") + "\n\n" + ocr_text
    elif "word" in content_type or "officedocument" in content_type or file_path.lower().endswith(".docx"):
        raw_text = _extract_docx_text(file_path)
    else:
        raise ValueError(f"Unsupported document format: {content_type}")

    raw_text = clean_text(raw_text)
    if not raw_text:
        raise ValueError("The uploaded document appears to be empty or contains unreadable text.")

    # 2. Extract contact info
    contact_info = _parse_contact_info(raw_text)

    # 3. Segment into sections
    parsed_sections = _segment_sections(raw_text, contact_info)

    return raw_text, parsed_sections


def _extract_pdf_text(file_path: str) -> str:
    """Extract raw text from a PDF file using pdfplumber."""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            else:
                logger.warning(f"No text extracted from page {page_idx + 1} of PDF.")
    return "\n\n".join(text_parts)


def _extract_pdf_text_with_ocr(file_path: str) -> str:
    """Fallback OCR extraction for PDFs using pdf2image + pytesseract when native text extraction fails."""
    if not _OCR_AVAILABLE:
        _LOGGER.warning("OCR dependencies not available (pdf2image/pytesseract). Skipping OCR fallback.")
        return ""

    text_parts = []
    try:
        images = convert_from_path(file_path, dpi=300)
        for img in images:
            try:
                text = pytesseract.image_to_string(img)
                if text and text.strip():
                    text_parts.append(text)
            except Exception as exc:
                _LOGGER.warning(f"Tesseract OCR failed on page image: {exc}")
    except Exception as exc:
        _LOGGER.warning(f"pdf2image failed to convert PDF to images for OCR: {exc}")

    return "\n\n".join(text_parts)


def _extract_docx_text(file_path: str) -> str:
    """Extract raw text from a DOCX file using python-docx."""
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # Process tables as well
    for table in doc.tables:
        # Attempt to detect column headers to handle project tables
        header_row = table.rows[0]
        header_cells = [re.sub(r"[^a-z0-9 ]", "", c.text.strip().lower()) for c in header_row.cells]
        # keywords to identify project-style tables
        header_keywords = [
            "project",
            "project name",
            "title",
            "description",
            "summary",
            "tech",
            "technology",
            "stack",
            "tools",
            "repository",
            "link",
            "url",
        ]

        # fuzzy detection: if any header cell contains any keyword
        is_project_table = any(any(k in h for k in header_keywords) for h in header_cells)

        # build header->column index map when possible
        col_map = {}
        for idx, h in enumerate(header_cells):
            for k in header_keywords:
                if k in h:
                    col_map[k] = idx

        for r_idx, row in enumerate(table.rows):
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if not row_texts:
                continue
            if is_project_table and len(table.rows) > 1 and r_idx > 0:
                cols = [c.text.strip() for c in row.cells]
                # heuristically find name/desc/tech/link using header map or heuristic content
                name = ""
                desc = ""
                tech = ""
                link = ""

                # try header-driven mapping
                if any(k in col_map for k in ["project name", "project", "title"]):
                    idx_name = col_map.get("project name") or col_map.get("project") or col_map.get("title")
                    name = cols[idx_name].strip() if idx_name is not None and idx_name < len(cols) else ""
                if any(k in col_map for k in ["description", "summary"]):
                    idx_desc = col_map.get("description") or col_map.get("summary")
                    desc = cols[idx_desc].strip() if idx_desc is not None and idx_desc < len(cols) else ""
                if any(k in col_map for k in ["tech", "technology", "stack", "tools"]):
                    idx_tech = col_map.get("tech") or col_map.get("technology") or col_map.get("stack") or col_map.get("tools")
                    tech = cols[idx_tech].strip() if idx_tech is not None and idx_tech < len(cols) else ""
                if any(k in col_map for k in ["repository", "link", "url"]):
                    idx_link = col_map.get("repository") or col_map.get("link") or col_map.get("url")
                    link = cols[idx_link].strip() if idx_link is not None and idx_link < len(cols) else ""

                # heuristics if header mapping didn't yield values
                if not name:
                    # choose the longest non-url cell as name candidate if short
                    candidates = [(i, c) for i, c in enumerate(cols) if c and not re.search(r"http|github.com|bitbucket.org", c.lower())]
                    if candidates:
                        # pick first short cell as name
                        name = candidates[0][1]
                if not link:
                    for c in cols:
                        if c.startswith("http") or "github.com" in c.lower() or "bitbucket.org" in c.lower():
                            link = c
                            break
                if not tech:
                    # infer tech from any cell containing common tech separators or keywords
                    tech_candidates = []
                    for c in cols:
                        if any(tok in c.lower() for tok in [",", ";", "python", "react", "django", "flask", "tensorflow", "pytorch", "docker", "kubernetes", "aws", "postgres", "sql", "mongodb"]):
                            tech_candidates.append(c)
                    tech = tech_candidates[0] if tech_candidates else tech

                proj_line = f"Project: {name} - {desc} - Tech: {tech}"
                if link:
                    proj_line += f" - Link: {link}"
                paragraphs.append(proj_line)
            else:
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

    return "\n".join(paragraphs)


def _parse_contact_info(text: str) -> Dict[str, Any]:
    """Parse contact info like Name, Email, Phone, URLs from the first few lines of text."""
    lines = [line.strip() for line in text.split("\n") if line.strip()][:10]
    
    # Simple name heuristic: first non-empty line that doesn't contain contact symbols
    name = "Candidate Name"
    for line in lines:
        if len(line) < 50 and not any(sym in line for sym in ["@", "+", "www.", "http", "|", "/", "\\", "github", "linkedin"]):
            # Confirm it's mostly words
            if re.match(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+$", line):
                name = line
                break
            elif re.match(r"^[a-zA-Z\s]+$", line) and len(line.split()) >= 2:
                name = line
                break

    emails = extract_emails(text)
    phones = extract_phones(text)
    urls = extract_urls(text)

    email = emails[0] if emails else ""
    phone = phones[0] if phones else ""
    linkedin = next((url for url in urls if "linkedin.com" in url.lower()), "")
    github = next((url for url in urls if "github.com" in url.lower()), "")
    portfolio = next((url for url in urls if url != linkedin and url != github), "")

    # Location heuristic: look for "City, State" or "City, Country" in lines near contact details
    location = ""
    loc_pattern = r"\b[A-Z][a-zA-Z\s\.\-]+,\s*[A-Z]{2,}\b|\b[A-Z][a-zA-Z\s\.\-]+,\s*[A-Z][a-zA-Z\s\.\-]+\b"
    for line in lines[:8]:
        match = re.search(loc_pattern, line)
        if match:
            location = match.group(0)
            break

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
        "location": location,
    }


def _segment_sections(text: str, contact_info: Dict[str, Any]) -> Dict[str, Any]:
    """Segment resume text into standard sections based on keywords and header rules."""
    lines = text.split("\n")
    
    sections: Dict[str, list] = {k: [] for k in SECTION_PATTERNS.keys()}
    sections["header"] = []
    current_section = "header"  # start in "header" section
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Detect section boundary headers (typically short lines with specific keywords)
        is_header = False
        if len(line_stripped) < 40:
            # Check matches across patterns
            for section_name, pattern in SECTION_PATTERNS.items():
                if re.match(pattern, line_stripped.lower()):
                    current_section = section_name
                    is_header = True
                    break
        
        if not is_header:
            sections[current_section].append(line)

    # Convert sections to output formats matching ParsedSections schema
    raw_experience = _parse_experience(sections["experience"])
    education_entries = _parse_education(sections["education"])
    projects_entries = _parse_projects(sections["projects"])
    
    # Auto-routing projects listed under experience to the projects section
    experience_entries = []
    for entry in raw_experience:
        role = entry.get("role", "")
        company = entry.get("company", "")
        bullets = entry.get("bullets", [])
        
        # Check if this is actually a project
        is_project = False
        if any(kw in role.lower() or kw in company.lower() for kw in ["project", "airbnb", "prediction", "dataset", "classifier", "model", "segmentation", "recognition", "detection", "neural network", "regression", "classification"]):
            is_project = True
            
        if is_project:
            # Extract clean project name
            p_name = company
            if p_name == "Various Companies" or len(p_name) < 4 or any(kw in p_name.lower() for kw in ["researcher", "engineer", "developer", "intern"]):
                p_name = role
                
            if "project" in p_name.lower():
                p_name = re.sub(r"^.*?project\s*", "", p_name, flags=re.IGNORECASE)
            if "researcher" in p_name.lower():
                p_name = re.sub(r"^.*?researcher\s*", "", p_name, flags=re.IGNORECASE)
                
            p_name = re.sub(r"^[\s\-–—,|]+|[\s\-–—,|]+$", "", p_name).strip()
            if not p_name:
                p_name = role
                
            projects_entries.append({
                "name": p_name,
                "technologies": [],
                "description": bullets,
                "link": ""
            })
        else:
            experience_entries.append(entry)
    
    skills_raw = "\n".join(sections["skills"])
    skills_list = [s.strip() for s in re.split(r"[,;•\n|]+", skills_raw) if s.strip() and len(s.strip()) < 40]

    cert_raw = "\n".join(sections["certifications"])
    cert_list = [c.strip() for c in re.split(r"[,;•\n]+", cert_raw) if c.strip() and len(c.strip()) < 80]

    awards_raw = "\n".join(sections["awards"])
    awards_list = [a.strip() for a in re.split(r"[,;•\n]+", awards_raw) if a.strip()]

    lang_raw = "\n".join(sections["languages"])
    lang_list = [l.strip() for l in re.split(r"[,;•\n]+", lang_raw) if l.strip() and len(l.strip()) < 30]

    pub_raw = "\n".join(sections["publications"])
    pub_list = [p.strip() for p in re.split(r"\n+[-•]?\s*", pub_raw) if p.strip()]

    # Volunteer
    volunteer_entries = []
    if sections["volunteer"]:
        volunteer_entries = [{"description": "\n".join(sections["volunteer"])}]

    # Summary text resolution:
    # If summary section is empty, heuristically extract actual summary prose from the header block
    summary_text = "\n".join(sections["summary"]).strip()
    if not summary_text and sections["header"]:
        header_prose = []
        name_lower = contact_info.get("name", "").lower()
        email_lower = contact_info.get("email", "").lower()
        phone = contact_info.get("phone", "")
        location_lower = contact_info.get("location", "").lower()
        
        for line in sections["header"]:
            l_strip = line.strip()
            l_lower = l_strip.lower()
            # Skip if empty or short
            if not l_strip or len(l_strip) < 3:
                continue
            # Skip if it is candidate name (case-insensitive name substring check)
            if name_lower and (name_lower in l_lower or l_lower in name_lower):
                continue
            # Skip if it contains email or is email
            if email_lower and email_lower in l_lower:
                continue
            # Skip if it contains phone
            if phone and phone in l_lower:
                continue
            # Skip if it contains location
            if location_lower and location_lower in l_lower:
                continue
            # Skip if it contains other contact signatures
            if any(sig in l_lower for sig in ["github.com", "linkedin.com", "@", "www.", "http"]):
                continue
            # Keep anything else that looks like summary prose
            header_prose.append(line)
        summary_text = "\n".join(header_prose).strip()

    return {
        "contact_info": contact_info,
        "summary": summary_text,
        "experience": experience_entries,
        "education": education_entries,
        "skills": skills_list,
        "certifications": cert_list,
        "projects": projects_entries,
        "languages": lang_list,
        "awards": awards_list,
        "publications": pub_list,
        "volunteer": volunteer_entries,
        "additional": "",
    }


def clean_role_company(role: str, company: str) -> tuple[str, str]:
    role = role.strip()
    company = company.strip()
    
    # 1. Resolve merge contamination bugs (e.g. company is "Various Companies" but role contains delimiter)
    if company.lower() in ["various companies", "various", ""] or not company:
        # Check if role contains a delimiter like ",", "@", "at", "—"
        for sep in [r"\s+at\s+", r"\s+@\s+", r"\s*—\s*", r"\s*–\s*", r"\s*-\s*", r"\s*,\s*"]:
            parts = re.split(sep, role, maxsplit=1, flags=re.IGNORECASE)
            if len(parts) == 2:
                potential_role = parts[0].strip()
                potential_company = parts[1].strip()
                if potential_company and len(potential_company) > 2 and not any(k in potential_company.lower() for k in ["bengaluru", "karnataka", "india", "bahama"]):
                    role = potential_role
                    company = potential_company
                    break
                    
    role_keywords = [
        "lead", "engineer", "developer", "analyst", "intern", "manager", "specialist", 
        "researcher", "consultant", "architect", "designer", "programmer", "head", 
        "director", "coordinator", "officer", "administrator"
    ]
    role_lower = role.lower()
    company_lower = company.lower()
    
    # Swap check: if first part is very short or looks like a company name and second has role keywords
    has_role_in_company = any(re.search(rf"\b{kw}\b", company_lower) for kw in role_keywords)
    has_role_in_role = any(re.search(rf"\b{kw}\b", role_lower) for kw in role_keywords)
    
    if has_role_in_company and not has_role_in_role:
        role, company = company, role
        role_lower, company_lower = role.lower(), company.lower()
        has_role_in_role = True
        has_role_in_company = False
        
    # De-duplication of role name inside company name
    if role_lower and company_lower:
        company = re.sub(rf"^\b{re.escape(role_lower)}\b[\s\-–—,|]*", "", company, flags=re.IGNORECASE).strip()
        company = re.sub(rf"[\s\-–—,|]*\b{re.escape(role_lower)}\b$", "", company, flags=re.IGNORECASE).strip()
        if role_lower in company_lower:
            company = company.replace(role, "").strip()
            company = company.replace(role.title(), "").strip()
            company = company.replace(role.upper(), "").strip()
            company = company.replace(role.lower(), "").strip()
            
        company = re.sub(r"^[\s\-–—,|]+|[\s\-–—,|]+$", "", company).strip()
        company_lower = company.lower()
        
    # E.g. Role: "Intern", Company: "Data Analyst Intern Company B" -> Role: "Data Analyst Intern", Company: "Company B"
    if role_lower == "intern" and "intern" in company_lower:
        match = re.search(r"\b([a-zA-Z\s]+intern)\b", company_lower)
        if match:
            matched_role = match.group(1).title()
            comp_clean = re.sub(re.escape(match.group(1)), "", company_lower).strip()
            idx = company_lower.find(comp_clean)
            if idx != -1:
                company = company[idx:idx+len(comp_clean)].strip()
            else:
                company = comp_clean.title()
            role = matched_role
            
    # Clean leading/trailing punctuation
    role = re.sub(r"^[\s\-–—,|]+|[\s\-–—,|]+$", "", role).strip()
    company = re.sub(r"^[\s\-–—,|]+|[\s\-–—,|]+$", "", company).strip()
    
    if not role:
        role = "Software Engineer"
    if not company:
        company = "Various Companies"
        
    return role, company


def _parse_experience(lines: list[str]) -> list[Dict[str, Any]]:
    """Helper to convert experience text lines into structured entries using smart sequential boundary detection."""
    if not lines:
        return []
        
    cleaned_lines = [l.strip() for l in lines if l.strip()]
    if not cleaned_lines:
        return []
        
    # 1. Split lines into distinct job blocks using sequential scanning
    blocks = []
    current_block = []
    
    role_keywords = [
        "engineer", "developer", "lead", "architect", "analyst", "intern", "manager", 
        "specialist", "researcher", "consultant", "director", "coordinator", "officer", 
        "administrator", "process engineer", "lead developer", "lead architect"
    ]
    
    for line in cleaned_lines:
        is_boundary = False
        is_bullet = bool(re.match(r"^[\s\-•\*·]+\s*", line))
        
        if not is_bullet and len(line) < 120:
            has_dates = bool(re.search(r"\b\d{4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\b", line, re.IGNORECASE))
            has_dash = any(sep in line for sep in ["—", "–", "―", "|"])
            line_lower = line.lower()
            has_role_kw = any(re.search(rf"\b{kw}\b", line_lower) for kw in role_keywords)
            
            if (has_dash and has_dates) or (has_dash and has_role_kw) or (has_dates and has_role_kw) or ("company" in line_lower) or ("project" in line_lower):
                is_boundary = True
                
        if is_boundary and current_block:
            blocks.append(current_block)
            current_block = []
            
        current_block.append(line)
        
    if current_block:
        blocks.append(current_block)
        
    # 2. Parse each block into a structured entry
    entries = []
    for block_lines in blocks:
        if not block_lines:
            continue
            
        header = block_lines[0].strip()
        bullets_start_idx = 1
        
        # Heuristic for multi-line headers: role and company split across two lines
        if len(block_lines) > 1:
            second_line = block_lines[1].strip()
            is_bullet = bool(re.match(r"^[\s\-•\*·]+\s*", second_line))
            if not is_bullet and len(second_line) < 60:
                role_keywords = ["engineer", "developer", "lead", "architect", "analyst", "intern", "manager", "specialist", "researcher", "consultant", "director", "coordinator", "officer", "administrator"]
                has_role_first = any(kw in header.lower() for kw in role_keywords)
                has_role_second = any(kw in second_line.lower() for kw in role_keywords)
                
                if has_role_second and not has_role_first:
                    header = f"{second_line} — {header}"
                    bullets_start_idx = 2
                elif has_role_first and not has_role_second:
                    header = f"{header} — {second_line}"
                    bullets_start_idx = 2
                elif not has_role_first and not has_role_second:
                    if any(loc in header.lower() for loc in ["company", "inc", "ltd", "corp", "university", "bahama", "karnataka", "bengaluru"]):
                        header = f"{second_line} — {header}"
                        bullets_start_idx = 2

        bullets = []
        for l in block_lines[bullets_start_idx:]:
            line_clean = l.strip()
            if not line_clean:
                continue
                
            is_new_bullet = False
            # Starts with bullet marker
            if re.match(r"^(?:-+|•+|\*+|·+|\(cid:\d+\))\s*", line_clean):
                is_new_bullet = True
            elif not bullets:
                is_new_bullet = True
            else:
                prev_bullet = bullets[-1].strip()
                if prev_bullet and prev_bullet[-1] in [".", "!", "?"]:
                    is_new_bullet = True
                    
            clean_b = re.sub(r"^(?:-+|•+|\*+|·+|\(cid:\d+\))\s*", "", line_clean)
            clean_b = re.sub(r"\(cid:\d+\)", "", clean_b).strip()
            
            if is_new_bullet:
                bullets.append(clean_b)
            else:
                bullets[-1] = heal_bullet_text(bullets[-1], clean_b)
                
        # Parse header for dates (like 2020 - 2022, or Jan 2021 - Present)
        date_pattern = r"[\(\[\s]*\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December|\d{1,2}/\d{2,4})\s*\d{4}\s*[-–—]\s*(?:Present|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December|\d{1,2}/\d{2,4})\s*\d{4}|\d{4})\b[\)\]\s]*|[\(\[\s]*\b\d{4}\s*[-–—]\s*(?:\d{4}|Present)\b[\)\]\s]*"
        dates = ""
        date_match = re.search(date_pattern, header)
        if date_match:
            dates = date_match.group(0)
            header = header.replace(dates, "").strip()
            
        # Remove empty parentheses/brackets left over from date extraction
        header = re.sub(r"\(\s*\)", "", header).strip()
        header = re.sub(r"\[\s*\]", "", header).strip()
        # Clean trailing separators
        header = re.sub(r"[\s\-–—,|]+$", "", header).strip()
 
        # Parse role and company from header using all delimiters
        delimiters = [r"\|", r"–", r"—", r"―", r"\s+-\s+", r"\bat\b", r"\bfor\b", r"\b@\b"]
        delims_regex = "|".join(delimiters)
        parts = re.split(delims_regex, header, maxsplit=1, flags=re.IGNORECASE)
        
        if len(parts) == 2:
            role = parts[0].strip()
            company = parts[1].strip()
        else:
            role = header.strip()
            company = "Various Companies"
            
        # Run smart role-company de-duplication and swap cleanup heuristics
        role, company = clean_role_company(role, company)
 
        if not bullets and len(block_lines) > bullets_start_idx:
            # If bullet points aren't prefixed, just treat them as list entries
            bullets = [line.strip() for line in block_lines[bullets_start_idx:]]
 
        entries.append({
            "company": company,
            "role": role,
            "dates": re.sub(r"^[\(\[\s,]+|[\)\]\s,]+$", "", dates).strip(), # clean dates punctuation
            "bullets": bullets if bullets else ["Performed daily duties as part of the core operational team."]
        })

    return entries


def _parse_education(lines: list[str]) -> list[Dict[str, Any]]:
    """Helper to convert education lines into structured entries."""
    text = "\n".join(lines)
    if not text.strip():
        return []
    
    blocks = re.split(r"\n\n+", text)
    entries = []
    
    for block in blocks:
        block_lines = [l.strip() for l in block.split("\n") if l.strip()]
        if not block_lines:
            continue
        
        header = block_lines[0]
        # Look for dates
        date_match = re.search(r"\b\d{4}(?:\s*-\s*(?:\d{4}|Present))?\b", header)
        dates = date_match.group(0) if date_match else ""
        if dates:
            header = header.replace(dates, "").strip()

        degree = header
        school = "Accredited University"
        for sep in ["|", "-", "at", ",", "from"]:
            if sep in header:
                parts = header.split(sep, 1)
                degree = parts[0].strip()
                school = parts[1].strip()
                break

        gpa = ""
        gpa_match = re.search(r"\bGPA:?\s*(\d\.\d{1,2}(?:/\d\.\d)?)\b", block.upper())
        if gpa_match:
            gpa = gpa_match.group(1)

        entries.append({
            "school": school,
            "degree": degree,
            "dates": dates,
            "gpa": gpa,
        })
    
    return entries


def _parse_projects(lines: list[str]) -> list[Dict[str, Any]]:
    """Helper to convert projects text into structured entries with smart boundary detection."""
    if not lines:
        return []
        
    cleaned_lines = [l.strip() for l in lines if l.strip()]
    if not cleaned_lines:
        return []
        
    project_blocks = []
    current_block = []
    
    # Smart boundary detection
    for idx, line in enumerate(cleaned_lines):
        is_boundary = False
        
        # Rule 1: Starts with explicit project keyword
        if re.match(r"^project\b", line, re.IGNORECASE) and not re.search(r"bullet|description|detail", line, re.IGNORECASE):
            is_boundary = True
            
        # Rule 2: Short line and next line starts with Tech Stack or bullet points, and we already have some lines in current block
        elif current_block and len(line) < 50 and not line.startswith(("-", "•", "*", "•")):
            # Check if next line contains Tech Stack / Technologies / Tools / Built with
            next_line = cleaned_lines[idx + 1] if idx + 1 < len(cleaned_lines) else ""
            if next_line and (
                re.search(r"\b(technologies|tech stack|built with|tools?):\s*", next_line, re.IGNORECASE) or
                next_line.startswith(("-", "•", "*", "•"))
            ):
                is_boundary = True
                
        # Rule 3: Contains a GitHub URL at the start or on a short line
        elif "github.com/" in line.lower() and len(line) < 80 and current_block:
            # If the current block already has description bullets, this github URL is likely the start of a new project header
            has_bullets = any(l.startswith(("-", "•", "*")) for l in current_block)
            if has_bullets:
                is_boundary = True
                
        if is_boundary and current_block:
            project_blocks.append(current_block)
            current_block = []
            
        current_block.append(line)
        
    if current_block:
        project_blocks.append(current_block)
        
    entries = []
    for block_lines in project_blocks:
        if not block_lines:
            continue
            
        # Parse the block
        name = block_lines[0]
        # Clean prefix like "Project:", "Personal Project:"
        name = re.sub(r"^(project|personal project|academic project|selected project):\s*", "", name, flags=re.IGNORECASE)
        # Clean trailing separators
        name = re.sub(r"[\s\-\|]+$", "", name).strip()
        # Remove (cid:X) tags
        name = re.sub(r"\(cid:\d+\)", "", name).strip()
        
        # Handle single-line projects
        if len(block_lines) == 1:
            parts = re.split(r"\s+[-|]\s+", name)
            if len(parts) > 1:
                name = parts[0].strip()
                block_lines.extend([p.strip() for p in parts[1:]])
        
        # Extract GitHub link
        link = ""
        for line in block_lines:
            if "github.com/" in line.lower() or "http" in line.lower():
                urls = extract_urls(line)
                if urls:
                    link = urls[0]
                    break
                    
        # Extract technologies and build clean description bullets
        technologies = []
        bullets = []
        
        for line in block_lines[1:]:
            line_cleaned = re.sub(r"\(cid:\d+\)", "", line).strip()
            if not line_cleaned:
                continue
                
            # Check if this line is tech stack
            tech_match = re.search(r"\b(tech|technologies|tech stack|built with|tools?):\s*(.*)\b", line_cleaned, re.IGNORECASE)
            if tech_match:
                tech_vals = tech_match.group(2)
                technologies = [t.strip() for t in re.split(r"[,;•|]+", tech_vals) if t.strip()]
                continue
                
            # If it's a URL and we already saved it, skip or don't put in bullets
            if link and link in line_cleaned and len(line_cleaned) < 80:
                continue
                
            # Clean bullets prefix
            clean_b = re.sub(r"^[\s\-•\*]+\s*", "", line_cleaned).strip()
            if clean_b:
                bullets.append(clean_b)
                
        # Title cleanups via normalize_project_name from github_service
        try:
            from app.services.github_service import normalize_project_name
            name = normalize_project_name(name)
        except Exception:
            pass
            
        entries.append({
            "name": name,
            "technologies": technologies,
            "description": bullets,
            "link": link
        })
        
    return entries
