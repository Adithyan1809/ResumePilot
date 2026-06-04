"""
Resume document generation service for PDF and DOCX formats.

Generates beautiful:
 PDF resumes via WeasyPrint + Jinja2 HTML templates
 DOCX resumes via docxtpl + Word templates
- DOCX resumes via docxtpl + Word templates
Includes a highly robust ReportLab PDF fallback if WeasyPrint dependencies
(Cairo/Pango libraries) are missing in the local execution environment.
"""

import logging
import re
from pathlib import Path
from typing import Any, Dict
from xml.sax.saxutils import escape

from app.core.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()

# Setup Jinja2 template loader
TEMPLATES_DIR = Path(__file__).parents[2] / "templates"


def clean_rendered_field(val: Any) -> str:
    """Helper to clean and strip empty placeholders, double spaces, and formatting artifacts from fields."""
    if not val:
        return ""
    if isinstance(val, dict):
        val = val.get("text", "")
    val_str = str(val).strip()
    val_str = re.sub(r"\(\s*\)", "", val_str)
    val_str = re.sub(r"-\s*$", "", val_str)
    val_str = re.sub(r"^\s*-", "", val_str)
    val_str = re.sub(r",\s*$", "", val_str)
    val_str = re.sub(r"\s+", " ", val_str)
    return val_str.strip()


def unwrap_text_field(value: Any) -> str:
    """Convert structured text field objects into plain strings."""
    if isinstance(value, dict):
        return str(value.get("text", "")).strip()
    return str(value).strip() if value is not None else ""


def normalize_resume_sections(sections: Dict[str, Any]) -> Dict[str, Any]:
    """Flatten structured section values into render-friendly strings/lists."""
    normalized = {k: v for k, v in sections.items()}

    summary = normalized.get("summary", "")
    normalized["summary"] = unwrap_text_field(summary)

    experience_items = []
    for entry in normalized.get("experience", []) or []:
        entry_copy = dict(entry) if isinstance(entry, dict) else {"role": unwrap_text_field(entry)}
        entry_copy["role"] = unwrap_text_field(entry_copy.get("role", ""))
        entry_copy["company"] = unwrap_text_field(entry_copy.get("company", ""))
        entry_copy["dates"] = unwrap_text_field(entry_copy.get("dates", ""))
        entry_copy["bullets"] = [unwrap_text_field(b) for b in (entry_copy.get("bullets", []) or [])]
        experience_items.append(entry_copy)
    normalized["experience"] = experience_items

    skills_items = []
    for skill in normalized.get("skills", []) or []:
        text = unwrap_text_field(skill)
        if text:
            skills_items.append(text)
    normalized["skills"] = skills_items

    education_items = []
    for entry in normalized.get("education", []) or []:
        entry_copy = dict(entry) if isinstance(entry, dict) else {"degree": unwrap_text_field(entry)}
        entry_copy["degree"] = unwrap_text_field(entry_copy.get("degree", ""))
        entry_copy["school"] = unwrap_text_field(entry_copy.get("school", ""))
        entry_copy["dates"] = unwrap_text_field(entry_copy.get("dates", ""))
        entry_copy["gpa"] = unwrap_text_field(entry_copy.get("gpa", ""))
        education_items.append(entry_copy)
    normalized["education"] = education_items

    return normalized


def normalize_skills(skills: list[str]) -> list[str]:
    """Helper to group flat skills list into standardized, normalized categories and deduplicate them.
    
    Standardized Categories:
    - Programming Languages
    - Data Science & ML
    - Backend Development
    - Databases
    - Computer Vision
    - DevOps & Infrastructure
    - Tools
    """
    import re
    
    # 7 standard categories
    standard_categories = [
        "Programming Languages",
        "Data Science & ML",
        "Backend Development",
        "Databases",
        "Computer Vision",
        "DevOps & Infrastructure",
        "Tools"
    ]
    
    categorized = {cat: [] for cat in standard_categories}
    
    # Skill vocabulary for classification
    skill_mapping = {
        "Programming Languages": [
            "python", "javascript", "typescript", "c++", "c#", "java", "html", "css", "go", "rust", "ruby", "php", 
            "sql", "shell", "bash", "r", "kotlin", "swift", "scala", "c", "assembly", "perl"
        ],
        "Data Science & ML": [
            "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "keras", "machine learning", "deep learning", 
            "nlp", "natural language processing", "data analysis", "statistics", "scipy", "matplotlib", "seaborn", 
            "hugging face", "llm", "bert", "gpt", "analytics", "data science", "xgboost", "lightgbm"
        ],
        "Backend Development": [
            "fastapi", "flask", "django", "node.js", "node", "express", "rest api", "rest apis", "graphql", 
            "websockets", "grpc", "microservices", "api gateway", "celery", "backend", "mvc", "apis", "spring boot"
        ],
        "Databases": [
            "postgresql", "sqlite", "mysql", "redis", "mongodb", "cassandra", "dynamodb", "oracle", "sql server", 
            "neo4j", "elasticsearch", "dbms", "caching", "nosql", "rdbms", "postgres"
        ],
        "Computer Vision": [
            "opencv", "yolo", "image processing", "cnn", "object detection", "pytesseract", "image segmentation", 
            "pil", "pillow", "gans", "resnet", "computer vision", "vision"
        ],
        "DevOps & Infrastructure": [
            "docker", "kubernetes", "aws", "gcp", "azure", "ci/cd", "github actions", "jenkins", "terraform", 
            "ansible", "linux", "nginx", "apache", "cloud", "serverless", "k8s"
        ],
        "Tools": [
            "git", "github", "gitlab", "postman", "jira", "confluence", "vs code", "webpack", "babel", "vite", 
            "npm", "yarn", "pip", "uv", "swagger", "windows", "macos"
        ]
    }
    
    # Track assigned skills to prevent cross-category duplication
    assigned_skills = set()
    raw_skills = []
    
    # 1. Parse all skills (flat or colon-grouped), recursively stripping duplicates and quotes
    for entry in skills:
        if not entry:
            continue
            
        if isinstance(entry, dict):
            entry_str = str(entry.get("text", "")).strip()
        else:
            entry_str = str(entry).strip()
            
        if not entry_str:
            continue
            
        # Clean quotes
        entry_str = entry_str.replace("'", "").replace('"', '').strip()
        
        # Split recursively by nested categories (e.g. "Programming Languages: Databases: Redis")
        while ":" in entry_str:
            parts = entry_str.split(":", 1)
            header = parts[0].strip()
            found = False
            for cat in standard_categories:
                if header.lower() == cat.lower() or header.lower().replace("&", "and") == cat.lower().replace("&", "and"):
                    entry_str = parts[1].strip()
                    found = True
                    break
            if not found:
                break
                
        sub_items = [s.strip() for s in re.split(r"[,;•|]+", entry_str) if s.strip()]
        raw_skills.extend(sub_items)
        
    # 2. Categorize each unique skill
    for skill in raw_skills:
        skill_clean = skill.strip()
        if not skill_clean or len(skill_clean) < 2:
            continue
            
        skill_lower = skill_clean.lower()
        if skill_lower in assigned_skills:
            continue
            
        matched_cat = None
        
        # Check standard lists
        for cat in ["Computer Vision", "Data Science & ML", "Backend Development", "Databases", "DevOps & Infrastructure", "Programming Languages"]:
            if any(kw == skill_lower or (len(kw) > 3 and kw in skill_lower) for kw in skill_mapping[cat]):
                matched_cat = cat
                break
                
        if not matched_cat:
            if any(kw == skill_lower or (len(kw) > 3 and kw in skill_lower) for kw in skill_mapping["Tools"]):
                matched_cat = "Tools"
                
        if not matched_cat:
            # Heuristics
            if "db" in skill_lower or "sql" in skill_lower:
                matched_cat = "Databases"
            elif any(w in skill_lower for w in ["docker", "k8s", "aws", "cloud", "pipeline"]):
                matched_cat = "DevOps & Infrastructure"
            elif any(w in skill_lower for w in ["api", "service", "web"]):
                matched_cat = "Backend Development"
            else:
                matched_cat = "Tools"
                
        categorized[matched_cat].append(skill_clean)
        assigned_skills.add(skill_lower)
        
    # 3. Format into category-colon strings with alphabetical sorting
    result = []
    for cat in standard_categories:
        skills_list = categorized[cat]
        if skills_list:
            # Sort alphabetically case-insensitively
            sorted_skills = sorted(list(set(skills_list)), key=lambda x: x.lower())
            result.append(f"{cat}: {', '.join(sorted_skills)}")
            
    return result


def _get_jinja_env():
    """Retrieve the Jinja2 environment configured for the templates directory."""
    from jinja2 import Environment, FileSystemLoader

    html_dir = TEMPLATES_DIR / "html"
    html_dir.mkdir(parents=True, exist_ok=True)
    return Environment(loader=FileSystemLoader(str(html_dir)))


async def generate_pdf(
    sections: Dict[str, Any],
    template_name: str,
    raw_text: str = "",
    original_sections: Dict[str, Any] = None,
    job_description: str = ""
) -> bytes:
    """Generate a print-ready PDF file from resume sections using a template.

    Args:
        sections: Tailored resume sections.
        template_name: The layout style (classic, modern, executive).
        raw_text: Optional original resume text for grounding checks.
        original_sections: Optional original resume parsed sections.
        job_description: Optional target job description.

    Returns:
        Bytes of the generated PDF file.
    """
    logger.info(f"Generating PDF using template '{template_name}'...")
    
    # ── Execute Strict Final Quality Gate ──
    from app.services.final_resume_quality_gate import validate_final_resume, QualityGateValidationError
    orig = original_sections or sections
    raw = raw_text or " ".join([
        sections.get("summary", ""),
        " ".join([b.get("text", "") if isinstance(b, dict) else str(b) for entry in sections.get("experience", []) for b in entry.get("bullets", [])]) if sections.get("experience") else "",
        " ".join(sections.get("skills", [])) if sections.get("skills") else ""
    ])
    
    gate_res = await validate_final_resume(
        sections=sections,
        raw_text=raw,
        original_sections=orig,
        job_description=job_description
    )
    sections = gate_res["sections"]
    
    # Dynamically compute layout if missing or empty
    if not sections.get("layout"):
        from app.services.content_density_manager import manage_content_density
        density_res = manage_content_density(sections)
        sections = {**sections, **density_res["sections"]}
        sections["layout"] = density_res["layout"]
        
    sections = normalize_resume_sections(sections)
    
    # ── WeasyPrint HTML Rendering ──────────────────────────────────
    try:
        from weasyprint import HTML, CSS
        
        env = _get_jinja_env()
        # Compile HTML string
        template_file = f"{template_name}.html"
        # Confirm it exists
        _ensure_html_templates_exist()
        
        # Normalize skills for perfect grouped rendering
        sections_copy = {k: v for k, v in sections.items()}
        sections_copy["skills"] = normalize_skills(sections.get("skills", []))
        
        template = env.get_template(template_file)
        html_out = template.render(
            resume=sections_copy,
            contact=sections_copy.get("contact_info", {}),
        )
        
        # Load style sheet
        css_path = TEMPLATES_DIR / "html" / "styles.css"
        
        # Compile PDF in-memory
        pdf_bytes = HTML(string=html_out).write_pdf(
            stylesheets=[CSS(str(css_path))] if css_path.exists() else None
        )
        logger.info("PDF generated successfully using WeasyPrint.")
        return pdf_bytes
        
    except Exception as exc:
        if isinstance(exc, QualityGateValidationError):
            raise
        logger.warning(
            f"WeasyPrint PDF generation failed: {exc}. "
            "Using robust ReportLab programmatic PDF layout engine fallback."
        )
        return await _generate_reportlab_pdf_fallback(sections, template_name, raw_text, original_sections, job_description)


async def generate_docx(
    sections: Dict[str, Any],
    template_name: str,
    raw_text: str = "",
    original_sections: Dict[str, Any] = None,
    job_description: str = ""
) -> bytes:
    """Generate a Microsoft Word (.docx) file from resume sections.

    Uses python-docx to construct an ATS-friendly single-column layout programmatically,
    providing high stability and complete freedom from corrupted template files.

    Args:
        sections: Tailored resume sections.
        template_name: The layout style (classic, modern, executive).
        raw_text: Optional original resume text for grounding checks.
        original_sections: Optional original resume parsed sections.
        job_description: Optional target job description.

    Returns:
        Bytes of the generated Word document.
    """
    logger.info(f"Generating DOCX using style '{template_name}'...")
    
    # ── Execute Strict Final Quality Gate ──
    from app.services.final_resume_quality_gate import validate_final_resume
    orig = original_sections or sections
    raw = raw_text or " ".join([
        sections.get("summary", ""),
        " ".join([b.get("text", "") if isinstance(b, dict) else str(b) for entry in sections.get("experience", []) for b in entry.get("bullets", [])]) if sections.get("experience") else "",
        " ".join(sections.get("skills", [])) if sections.get("skills") else ""
    ])
    
    gate_res = await validate_final_resume(
        sections=sections,
        raw_text=raw,
        original_sections=orig,
        job_description=job_description
    )
    sections = gate_res["sections"]
    
    # Dynamically compute layout if missing or empty
    if not sections.get("layout"):
        from app.services.content_density_manager import manage_content_density
        density_res = manage_content_density(sections)
        sections = {**sections, **density_res["sections"]}
        sections["layout"] = density_res["layout"]
        
    sections = normalize_resume_sections(sections)
    
    try:
        from io import BytesIO
        from docx import Document
        from docx.shared import Pt, Inches
        
        doc = Document()
        
        # Dynamically retrieve layout variables for DOCX spacing
        layout = sections.get("layout", {})
        margin_val = layout.get("page_margin_inches", 0.75)
        font_size_val = layout.get("font_size_pt", 10.5)
        space_after_section_val = layout.get("space_after_section", 8)
        
        body_font_size = font_size_val
        header_font_size = font_size_val + 1.5
        sub_font_size = font_size_val - 0.5
        tiny_font_size = font_size_val - 1.5
        title_font_size = font_size_val + 13.5
        
        # Configure margins dynamically based on content density
        sections_config = doc.sections
        for section in sections_config:
            section.top_margin = Inches(margin_val)
            section.bottom_margin = Inches(margin_val)
            section.left_margin = Inches(margin_val)
            section.right_margin = Inches(margin_val)

        # Define Styles & Fonts based on theme
        font_name = "Times New Roman" if template_name == "classic" else "Arial"
        
        # Add Contact Block
        contact = sections.get("contact_info", {})
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(contact.get("name", "Candidate Name"))
        title_run.font.name = font_name
        title_run.font.size = Pt(title_font_size)
        title_run.bold = True
        title_p.alignment = 1  # Center
        
        parts = [p.strip() for p in [contact.get('email', ''), contact.get('phone', ''), contact.get('location', '')] if p and p.strip()]
        contact_line = " | ".join(parts)
        contact_p = doc.add_paragraph()
        contact_run = contact_p.add_run(contact_line)
        contact_run.font.name = font_name
        contact_run.font.size = Pt(sub_font_size)
        contact_p.alignment = 1  # Center

        links = [contact.get("linkedin", ""), contact.get("github", "")]
        links = [l for l in links if l]
        if links:
            links_p = doc.add_paragraph()
            links_run = links_p.add_run(" | ".join(links))
            links_run.font.name = font_name
            links_run.font.size = Pt(tiny_font_size)
            links_p.alignment = 1  # Center

        doc.add_paragraph()  # spacing

        # Summary Section
        summary = sections.get("summary", "")
        if summary:
            _add_docx_section_header(doc, "PROFESSIONAL SUMMARY", font_name, font_size_pt=header_font_size, space_after_section_pt=space_after_section_val)
            p = doc.add_paragraph()
            run = p.add_run(summary)
            run.font.name = font_name
            run.font.size = Pt(body_font_size)

        # Experience Section
        experience = sections.get("experience", [])
        if experience:
            _add_docx_section_header(doc, "PROFESSIONAL EXPERIENCE", font_name, font_size_pt=header_font_size, space_after_section_pt=space_after_section_val)
            for entry in experience:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = int(space_after_section_val * 1.5 * 20)
                p.paragraph_format.space_after = int(space_after_section_val * 0.3 * 20)
                
                # Job Header
                role = clean_rendered_field(entry.get('role', 'Developer'))
                company = clean_rendered_field(entry.get('company', ''))
                dates = clean_rendered_field(entry.get('dates', ''))
                
                role_run = p.add_run(role)
                role_run.bold = True
                role_run.font.name = font_name
                role_run.font.size = Pt(body_font_size)
                
                if company:
                    p.add_run("  —  ")
                    company_run = p.add_run(company)
                    company_run.italic = True
                    company_run.font.name = font_name
                    company_run.font.size = Pt(body_font_size)

                if dates:
                    p.add_run("  ")
                    date_run = p.add_run(f"({dates})")
                    date_run.font.name = font_name
                    date_run.font.size = Pt(sub_font_size)

                # Bullets
                for bullet in entry.get("bullets", []):
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_after = int(space_after_section_val * 0.2 * 20)
                    bp_run = bp.add_run(bullet)
                    bp_run.font.name = font_name
                    bp_run.font.size = Pt(body_font_size)

        # Projects Section
        projects = sections.get("projects", [])
        if projects:
            _add_docx_section_header(doc, "PROJECTS", font_name, font_size_pt=header_font_size, space_after_section_pt=space_after_section_val)
            for entry in projects:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = int(space_after_section_val * 1.5 * 20)
                p.paragraph_format.space_after = int(space_after_section_val * 0.3 * 20)
                
                name = clean_rendered_field(entry.get('name', 'Project'))
                techs = entry.get('technologies', [])
                link = clean_rendered_field(entry.get('link', ''))
                
                name_run = p.add_run(name)
                name_run.bold = True
                name_run.font.name = font_name
                name_run.font.size = Pt(body_font_size)
                
                if techs:
                    p.add_run(" (")
                    tech_run = p.add_run(", ".join(techs))
                    tech_run.italic = True
                    tech_run.font.name = font_name
                    tech_run.font.size = Pt(sub_font_size)
                    p.add_run(")")
                
                if link:
                    p.add_run("  —  ")
                    link_run = p.add_run(link)
                    link_run.font.name = font_name
                    link_run.font.size = Pt(sub_font_size - 1.0)
                
                # Bullets
                description = entry.get("description", [])
                bullets = [description] if isinstance(description, str) else list(description) if isinstance(description, list) else []
                for bullet in bullets:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_after = int(space_after_section_val * 0.2 * 20)
                    bp_run = bp.add_run(bullet)
                    bp_run.font.name = font_name
                    bp_run.font.size = Pt(body_font_size)

        # Skills Section
        skills = normalize_skills(sections.get("skills", []))
        if skills:
            _add_docx_section_header(doc, "TECHNICAL SKILLS", font_name, font_size_pt=header_font_size, space_after_section_pt=space_after_section_val)
            for skill_entry in skills:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = int(space_after_section_val * 0.3 * 20)
                if ":" in skill_entry:
                    cat, items = skill_entry.split(":", 1)
                    
                    cat_run = p.add_run(f"{cat.strip()}: ")
                    cat_run.bold = True
                    cat_run.font.name = font_name
                    cat_run.font.size = Pt(body_font_size)
                    
                    items_run = p.add_run(items.strip())
                    items_run.font.name = font_name
                    items_run.font.size = Pt(body_font_size)
                else:
                    run = p.add_run(skill_entry.strip())
                    run.font.name = font_name
                    run.font.size = Pt(body_font_size)

        # Education Section
        education = sections.get("education", [])
        if education:
            _add_docx_section_header(doc, "EDUCATION", font_name, font_size_pt=header_font_size, space_after_section_pt=space_after_section_val)
            for entry in education:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = int(space_after_section_val * 0.3 * 20)
                deg = clean_rendered_field(entry.get('degree', 'Degree'))
                sch = clean_rendered_field(entry.get('school', ''))
                dates = clean_rendered_field(entry.get('dates', ''))
                gpa = clean_rendered_field(entry.get('gpa', ''))
                
                degree_run = p.add_run(deg)
                degree_run.bold = True
                degree_run.font.name = font_name
                degree_run.font.size = Pt(body_font_size)
                
                if sch:
                    p.add_run(" - ")
                    school_run = p.add_run(sch)
                    school_run.font.name = font_name
                    school_run.font.size = Pt(body_font_size)

                if dates:
                    p.add_run(" ")
                    date_run = p.add_run(f"({dates})")
                    date_run.font.name = font_name
                    date_run.font.size = Pt(sub_font_size)
                    
                if gpa:
                    p.add_run(f" (GPA: {gpa})")
                    # size run
                    for r in p.runs[-1:]:
                        r.font.name = font_name
                        r.font.size = Pt(body_font_size)
                        
                # Add Relevant Coursework support
                coursework = clean_rendered_field(entry.get('coursework', ''))
                if not coursework:
                    coursework = clean_rendered_field(entry.get('relevant_coursework', ''))
                if coursework:
                    cp = doc.add_paragraph()
                    cp.paragraph_format.space_before = int(space_after_section_val * 0.1 * 20)
                    cp.paragraph_format.space_after = int(space_after_section_val * 0.2 * 20)
                    c_run_label = cp.add_run("Relevant Coursework: ")
                    c_run_label.bold = True
                    c_run_label.font.name = font_name
                    c_run_label.font.size = Pt(sub_font_size)
                    c_run_val = cp.add_run(coursework)
                    c_run_val.font.name = font_name
                    c_run_val.font.size = Pt(sub_font_size)

        # Save to bytes stream
        stream = BytesIO()
        doc.save(stream)
        logger.info("DOCX document compiled successfully.")
        return stream.getvalue()

    except Exception as exc:
        logger.error(f"DOCX generation failed: {exc}")
        raise ValueError(f"Could not compile DOCX file: {exc}")


def _add_docx_section_header(doc, title: str, font_name: str, font_size_pt: float = 12.0, space_after_section_pt: float = 8.0):
    """Add a bold horizontal section header to a Word document."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = int(space_after_section_pt * 1.5 * 20)
    p.paragraph_format.space_after = int(space_after_section_pt * 0.5 * 20)
    run = p.add_run(title)
    run.bold = True
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    # Add a thin line under the title (border effect)
    p_format = p.paragraph_format
    p_format.keep_with_next = True


# ── Programmatic Fallback PDF Engine via ReportLab ──────────────

async def _generate_reportlab_pdf_fallback(
    sections: Dict[str, Any],
    template_name: str,
    raw_text: str = "",
    original_sections: Dict[str, Any] = None,
    job_description: str = ""
) -> bytes:
    """Fallback generator using standard ReportLab.

    ReportLab is purely Python-based and has no OS C-level library dependencies.
    """
    logger.info("Starting ReportLab programmatic PDF compilation...")
    
    # ── Execute Strict Final Quality Gate ──
    from app.services.final_resume_quality_gate import validate_final_resume
    orig = original_sections or sections
    raw = raw_text or " ".join([
        sections.get("summary", ""),
        " ".join([b.get("text", "") if isinstance(b, dict) else str(b) for entry in sections.get("experience", []) for b in entry.get("bullets", [])]) if sections.get("experience") else "",
        " ".join(sections.get("skills", [])) if sections.get("skills") else ""
    ])
    
    gate_res = await validate_final_resume(
        sections=sections,
        raw_text=raw,
        original_sections=orig,
        job_description=job_description
    )
    sections = gate_res["sections"]
    
    sections = normalize_resume_sections(sections)
    try:
        from io import BytesIO
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.platypus.flowables import HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
    except Exception as exc:
        logger.warning(f"ReportLab unavailable ({exc}). Using pure-Python PDF fallback.")
        return _generate_minimal_pdf_fallback(sections, template_name)

    buffer = BytesIO()
    
    # Dynamically retrieve layout variables for the PDF fallback generator
    layout = sections.get("layout", {})
    margin_in = layout.get("page_margin_inches", 0.75)
    margin_pt = margin_in * 72
    font_size_val = layout.get("font_size_pt", 10.5)
    line_spacing_val = layout.get("line_spacing", 1.2)
    leading_val = font_size_val * line_spacing_val
    space_after_section_val = layout.get("space_after_section", 8)

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=margin_pt,
        leftMargin=margin_pt,
        topMargin=margin_pt,
        bottomMargin=margin_pt,
    )
    
    # Set professional document metadata for browser tab titles
    candidate_name = sections.get("contact_info", {}).get("name", "Candidate")
    doc.title = f"{candidate_name} - Resume"
    doc.author = candidate_name
    doc.subject = f"Tailored Resume for {sections.get('job_title', 'Job Application')}"
    
    styles = getSampleStyleSheet()
    story = []
    
    # Configure font sizes/colors based on template style
    is_classic = template_name == "classic"
    primary_color = colors.HexColor("#0f172a") if is_classic else colors.HexColor("#1e1b4b")
    font_family = "Times-Roman" if is_classic else "Helvetica"
    font_bold = "Times-Bold" if is_classic else "Helvetica-Bold"

    # Define custom styles with dynamic dimensions
    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Normal"],
        fontName=font_bold,
        fontSize=font_size_val + 11.5,
        leading=font_size_val + 15.5,
        alignment=1,  # Center
        textColor=primary_color,
    )
    
    subtitle_style = ParagraphStyle(
        "DocSubtitle",
        parent=styles["Normal"],
        fontName=font_family,
        fontSize=font_size_val - 0.5,
        leading=font_size_val + 3.5,
        alignment=1,  # Center
        textColor=colors.HexColor("#475569"),
    )

    section_header_style = ParagraphStyle(
        "SectionHeader",
        parent=styles["Normal"],
        fontName=font_bold,
        fontSize=font_size_val + 1.5,
        leading=font_size_val + 5.5,
        textColor=primary_color,
        spaceBefore=space_after_section_val * 1.5,
        spaceAfter=space_after_section_val * 0.5,
        keepWithNext=True,
    )

    body_style = ParagraphStyle(
        "BodyTextCustom",
        parent=styles["Normal"],
        fontName=font_family,
        fontSize=font_size_val,
        leading=leading_val,
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=space_after_section_val * 0.5,
    )

    bullet_style = ParagraphStyle(
        "BulletCustom",
        parent=styles["Normal"],
        fontName=font_family,
        fontSize=font_size_val,
        leading=leading_val,
        textColor=colors.HexColor("#1e293b"),
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=space_after_section_val * 0.3,
    )

    # 1. Contact Block
    contact = sections.get("contact_info", {})
    story.append(Paragraph(escape(contact.get("name", "Candidate Name")), title_style))
    story.append(Spacer(1, 4))
    
    parts = [p.strip() for p in [contact.get('email', ''), contact.get('phone', ''), contact.get('location', '')] if p and p.strip()]
    details_line = " | ".join(parts)
    story.append(Paragraph(escape(details_line), subtitle_style))
    
    links = [contact.get("linkedin", ""), contact.get("github", "")]
    links = [l for l in links if l]
    if links:
        story.append(Paragraph(escape(" | ".join(links)), subtitle_style))
        
    story.append(Spacer(1, 10))

    # 2. Professional Summary
    summary = sections.get("summary", "")
    if summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_header_style))
        story.append(HRFlowable(width="100%", thickness=1, color=primary_color, spaceBefore=1, spaceAfter=8))
        story.append(Paragraph(escape(summary), body_style))
        story.append(Spacer(1, 6))

    # 3. Work Experience
    experience = sections.get("experience", [])
    if experience:
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", section_header_style))
        story.append(HRFlowable(width="100%", thickness=1, color=primary_color, spaceBefore=1, spaceAfter=8))
        
        for entry in experience:
            role = clean_rendered_field(entry.get("role", "Developer"))
            company = clean_rendered_field(entry.get("company", ""))
            dates = clean_rendered_field(entry.get("dates", ""))
            
            company_segment = f"  —  <i>{company}</i>" if company else ""
            date_segment = f"  ({dates})" if dates else ""
            hdr_text = f"<b>{escape(role)}</b>{company_segment}{date_segment}"
            story.append(Paragraph(hdr_text, body_style))
            
            for bullet in entry.get("bullets", []):
                story.append(Paragraph(escape(f"- {bullet}"), bullet_style))
            story.append(Spacer(1, 6))

    # Projects Section
    projects = sections.get("projects", [])
    if projects:
        story.append(Paragraph("PROJECTS", section_header_style))
        story.append(HRFlowable(width="100%", thickness=1, color=primary_color, spaceBefore=1, spaceAfter=8))
        
        for entry in projects:
            name = clean_rendered_field(entry.get("name", "Project"))
            techs = entry.get("technologies", [])
            link = clean_rendered_field(entry.get("link", ""))
            
            tech_segment = f" ({', '.join(techs)})" if techs else ""
            link_segment = f"  —  {link}" if link else ""
            hdr_text = f"<b>{escape(name)}</b><i>{escape(tech_segment)}</i>{escape(link_segment)}"
            story.append(Paragraph(hdr_text, body_style))
            
            description = entry.get("description", [])
            bullets = [description] if isinstance(description, str) else list(description) if isinstance(description, list) else []
            for bullet in bullets:
                story.append(Paragraph(escape(f"- {bullet}"), bullet_style))
            story.append(Spacer(1, 6))

    # 4. Technical Skills
    skills = normalize_skills(sections.get("skills", []))
    if skills:
        story.append(Paragraph("TECHNICAL SKILLS", section_header_style))
        story.append(HRFlowable(width="100%", thickness=1, color=primary_color, spaceBefore=1, spaceAfter=8))
        for skill_entry in skills:
            if ":" in skill_entry:
                cat, items = skill_entry.split(":", 1)
                skill_text = f"<b>{escape(cat.strip())}:</b> {escape(items.strip())}"
            else:
                skill_text = escape(skill_entry.strip())
            story.append(Paragraph(skill_text, body_style))
        story.append(Spacer(1, 6))

    # 5. Education
    education = sections.get("education", [])
    if education:
        story.append(Paragraph("EDUCATION", section_header_style))
        story.append(HRFlowable(width="100%", thickness=1, color=primary_color, spaceBefore=1, spaceAfter=8))
        for entry in education:
            deg = clean_rendered_field(entry.get("degree", "Degree"))
            sch = clean_rendered_field(entry.get("school", ""))
            dt = clean_rendered_field(entry.get("dates", ""))
            gpa = clean_rendered_field(entry.get("gpa", ""))
            
            school_segment = f", {sch}" if sch else ""
            date_segment = f" ({dt})" if dt else ""
            gpa_segment = f" (GPA: {gpa})" if gpa else ""
            edu_text = f"<b>{escape(deg)}</b>{escape(school_segment)}{escape(date_segment)}{escape(gpa_segment)}"
            story.append(Paragraph(edu_text, body_style))
            
            # Add Relevant Coursework support
            coursework = clean_rendered_field(entry.get("coursework", ""))
            if not coursework:
                coursework = clean_rendered_field(entry.get("relevant_coursework", ""))
            if coursework:
                cw_style = ParagraphStyle(
                    "CwCustom",
                    parent=styles["Normal"],
                    fontName=font_family,
                    fontSize=font_size_val - 0.5,
                    leading=leading_val - 1.0,
                    textColor=colors.HexColor("#475569"),
                    leftIndent=10,
                    spaceAfter=space_after_section_val * 0.3,
                )
                story.append(Paragraph(f"<b>Relevant Coursework:</b> {escape(coursework)}", cw_style))
            story.append(Spacer(1, 4))

    def set_metadata(canvas, doc_obj):
        canvas.setTitle(f"{candidate_name} - Resume")
        canvas.setAuthor(candidate_name)
        canvas.setSubject(f"Tailored Resume for {sections.get('job_title', 'Job Application')}")

    doc.build(story, onFirstPage=set_metadata, onLaterPages=set_metadata)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    
    logger.info("PDF generated successfully using ReportLab fallback.")
    return pdf_bytes


def _pdf_literal(text: str) -> str:
    """Escape text for a PDF literal string."""
    return (
        str(text)
        .replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
    )


def _generate_minimal_pdf_fallback(sections: Dict[str, Any], template_name: str) -> bytes:
    """Generate a tiny valid PDF without third-party dependencies."""
    sections = normalize_resume_sections(sections)
    contact = sections.get("contact_info", {}) or {}
    summary = sections.get("summary", "")
    experience = sections.get("experience", []) or []
    skills = sections.get("skills", []) or []
    education = sections.get("education", []) or []

    lines: list[str] = []
    name = contact.get("name", "Candidate Name")
    lines.append(name)
    contact_line = " | ".join(
        [x for x in [contact.get("email", ""), contact.get("phone", ""), contact.get("location", "")] if x]
    )
    if contact_line:
        lines.append(contact_line)
    if summary:
        lines.append("")
        lines.append("PROFESSIONAL SUMMARY")
        lines.append(summary)
    if experience:
        lines.append("")
        lines.append("PROFESSIONAL EXPERIENCE")
        for entry in experience:
            header = " - ".join([x for x in [entry.get("role", ""), entry.get("company", "")] if x])
            if entry.get("dates"):
                header = f"{header} ({entry.get('dates')})" if header else f"({entry.get('dates')})"
            lines.append(header)
            for bullet in entry.get("bullets", []):
                lines.append(f"- {bullet}")
    if skills:
        lines.append("")
        lines.append("TECHNICAL SKILLS")
        lines.extend(skills)
    if education:
        lines.append("")
        lines.append("EDUCATION")
        for entry in education:
            edu_bits = [entry.get("degree", ""), entry.get("school", ""), entry.get("dates", ""), entry.get("gpa", "")]
            lines.append(" | ".join([x for x in edu_bits if x]))

    if not lines:
        lines = ["ResumeAI", "Generated document"]

    # Keep the page content compact enough for a single page.
    lines = lines[:40]
    content_lines = ["BT", "/F1 11 Tf", "72 760 Td"]
    first = True
    for line in lines:
        if first:
            content_lines.append(f"({_pdf_literal(line)}) Tj")
            first = False
        else:
            content_lines.append("0 -14 Td")
            content_lines.append(f"({_pdf_literal(line)}) Tj")
    content_lines.append("ET")
    content_stream = "\n".join(content_lines).encode("latin-1", "replace")

    objects = []
    objects.append("1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n")
    objects.append("2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n")
    objects.append(
        "3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n"
    )
    objects.append("4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n")
    objects.append(f"5 0 obj << /Length {len(content_stream)} >> stream\n")
    objects.append(content_stream.decode("latin-1", "replace"))
    objects.append("\nendstream endobj\n")
    info_dict = f"<< /Title ({_pdf_literal(name)} - Resume) /Author ({_pdf_literal(name)}) >>"
    objects.append(f"6 0 obj {info_dict} endobj\n")

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj.encode("latin-1", "replace"))

    xref_start = len(pdf)
    pdf.extend(f"xref\n0 {len(offsets)}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        pdf.extend(f"{off:010d} 00000 n \n".encode("ascii"))
    pdf.extend(b"trailer << /Size 7 /Root 1 0 R /Info 6 0 R >>\nstartxref\n")
    pdf.extend(f"{xref_start}\n%%EOF".encode("ascii"))
    return bytes(pdf)


# ── Create Default Resume Output Templates ─────────────────────────

def _ensure_html_templates_exist():
    """Build placeholder HTML resume templates in backend/templates/html."""
    html_dir = TEMPLATES_DIR / "html"
    html_dir.mkdir(parents=True, exist_ok=True)
    
    # Styles.css (Always overwrite to ensure the new ATS style applies)
    css_file = html_dir / "styles.css"
    with open(css_file, "w", encoding="utf-8") as f:
        f.write("""
/* Print-optimized ATS Resume CSS Stylesheet */
body {
    font-family: 'Segoe UI', Arial, sans-serif;
    color: #000000;
    line-height: 1.35;
    font-size: 10pt;
    margin: 0;
    padding: 0;
}
.header {
    text-align: center;
    margin-bottom: 12px;
}
.header h1 {
    font-size: 20pt;
    font-weight: bold;
    color: #000000;
    text-transform: uppercase;
    letter-spacing: 1px;
    margin: 0 0 4px 0;
}
.contact-info {
    font-size: 9pt;
    color: #333333;
    margin-bottom: 2px;
}
.contact-info a {
    color: #5B4B8A;
    text-decoration: none;
}
.section-title {
    font-size: 10.5pt;
    font-weight: bold;
    color: #000000;
    margin-top: 14px;
    margin-bottom: 6px;
    border-bottom: 1px solid #000000;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    padding-bottom: 2px;
}
.summary {
    margin-bottom: 8px;
    text-align: justify;
}
.experience-item {
    margin-bottom: 8px;
}
.flex-row {
    display: flex;
    justify-content: space-between;
    align-items: baseline;
    margin-bottom: 2px;
}
.role-company b {
    font-weight: bold;
}
.dates {
    font-size: 9.5pt;
    color: #333333;
}
.sub-detail {
    font-style: italic;
    font-size: 9.5pt;
    color: #333333;
    margin-bottom: 4px;
}
.bullets {
    margin: 0;
    padding-left: 18px;
}
.bullets li {
    margin-bottom: 3px;
    text-align: justify;
}
.skills-list {
    margin-bottom: 8px;
}
.skills-row {
    margin-bottom: 3px;
}
.education-item {
    margin-bottom: 8px;
}
""")

    # HTML templates (Classic, Modern, Executive)
    for name in ["classic", "modern", "executive"]:
        html_file = html_dir / f"{name}.html"
        # Always overwrite in development to ensure new features render correctly
        with open(html_file, "w", encoding="utf-8") as f:
            f.write("""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{{ contact.name | default('Candidate') }} - Resume</title>
    <style>
        @page {
            size: letter;
            margin: {{ resume.layout.page_margin_inches | default(0.5) }}in;
        }
        body {
            font-size: {{ resume.layout.font_size_pt | default(10) }}pt;
            line-height: {{ resume.layout.line_spacing | default(1.35) }};
        }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ contact.name }}</h1>
        <div class="contact-info">
            {{ contact.location }}{% if contact.email %} | {{ contact.email }}{% endif %}{% if contact.phone %} | {{ contact.phone }}{% endif %}
        </div>
        <div class="contact-info">
            {% if contact.github %}<a href="{{ contact.github }}">{{ contact.github | replace('https://', '') | replace('http://', '') }}</a>{% endif %}
            {% if contact.linkedin %} | <a href="{{ contact.linkedin }}">{{ contact.linkedin | replace('https://', '') | replace('http://', '') }}</a>{% endif %}
            {% if contact.portfolio %} | <a href="{{ contact.portfolio }}">{{ contact.portfolio | replace('https://', '') | replace('http://', '') }}</a>{% endif %}
        </div>
    </div>

    {% if resume.summary %}
    <div class="section-title">Professional Summary</div>
    <div class="summary">
        {{ resume.summary }}
    </div>
    {% endif %}

    {% if resume.experience %}
    <div class="section-title">Work Experience</div>
    {% for job in resume.experience %}
    <div class="experience-item">
        <div class="flex-row">
            <div class="role-company"><b>{{ job.role }}</b>{% if job.company %} &mdash; {{ job.company }}{% endif %}</div>
            <div class="dates">{% if job.dates %}{{ job.dates }}{% endif %}</div>
        </div>
        {% if job.location %}
        <div class="sub-detail">{{ job.location }}</div>
        {% endif %}
        <ul class="bullets">
            {% for bullet in job.bullets %}
            <li>{{ bullet }}</li>
            {% endfor %}
        </ul>
    </div>
    {% endfor %}
    {% endif %}
    
    {% if resume.projects and resume.projects|length == 0 %}
    {% elif resume.projects %}
    <div class="section-title">Projects</div>
    {% for proj in resume.projects %}
    <div class="experience-item">
        <div class="flex-row">
            <div class="role-company"><b>{{ proj.name }}</b></div>
            <div class="dates">{% if proj.dates %}{{ proj.dates }}{% endif %}</div>
        </div>
        {% if proj.technologies or proj.link %}
        <div class="sub-detail">
            {% if proj.technologies %}{{ proj.technologies | join(' &middot; ') }}{% endif %}
            {% if proj.technologies and proj.link %} &middot; {% endif %}
            {% if proj.link %}<a href="{{ proj.link }}">Link</a>{% endif %}
        </div>
        {% endif %}
        <ul class="bullets">
            {% if proj.description is string %}
                <li>{{ proj.description }}</li>
            {% else %}
                {% for bullet in proj.description %}
                <li>{{ bullet }}</li>
                {% endfor %}
            {% endif %}
        </ul>
    </div>
    {% endfor %}
    {% endif %}

    {% if resume.education %}
    <div class="section-title">Education</div>
    {% for edu in resume.education %}
    <div class="education-item">
        <div class="flex-row">
            <div class="role-company"><b>{{ edu.degree }}</b></div>
            <div class="dates">{% if edu.dates %}{{ edu.dates }}{% endif %}</div>
        </div>
        <div class="sub-detail">
            {% if edu.school %}{{ edu.school }}{% endif %}
            {% if edu.school and edu.gpa %} | {% endif %}
            {% if edu.gpa %}CGPA: {{ edu.gpa }}{% endif %}
        </div>
        {% set coursework = edu.coursework | default(edu.relevant_coursework) %}
        {% if coursework %}
        <div class="sub-detail" style="margin-top: 2px;">
            <strong>Relevant Coursework:</strong> {{ coursework }}
        </div>
        {% endif %}
    </div>
    {% endfor %}
    {% endif %}

    {% if resume.skills %}
    <div class="section-title">Technical Skills</div>
    <div class="skills-list">
        {% for skill_entry in resume.skills %}
        <div class="skills-row">
            {% if ':' in skill_entry %}
                {% set parts = skill_entry.split(':', 1) %}
                <strong>{{ parts[0].strip() }}:</strong> {{ parts[1].strip() }}
            {% else %}
                {{ skill_entry }}
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if resume.certifications %}
    <div class="section-title">Certifications</div>
    <div class="skills-list">
        {% for cert in resume.certifications %}
        <div class="skills-row">
            {{ cert }}
        </div>
        {% endfor %}
    </div>
    {% endif %}
</body>
</html>
""")
